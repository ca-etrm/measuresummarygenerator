from __future__ import annotations
import os
import json
from docx.api import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.table import Table, _Row as Row
from docx.styles.style import ParagraphStyle, BaseStyle


def html_escape(html: str) -> str:
    return (
        html
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\"", "&quot;")
            .replace("\'", "&#39;")
    )


class RunInfo:
    def __init__(self, run: Run) -> None:
        self.bold = run.bold
        self.italic = run.italic
        self.underline = run.underline
        self.text = html_escape(run.text)

    @property
    def html(self) -> str:
        text = self.text
        if self.bold:
            text = f"<strong>{text}</strong>"

        if self.italic:
            text = f"<em>{text}</em>"

        if self.underline:
            text = f"<uline>{text}</uline>"

        return text

    def can_join(self, other: RunInfo) -> bool:
        if self.bold and not other.bold:
            return False

        if self.italic and not other.italic:
            return False

        if self.underline and not other.underline:
            return False

        return True

    def join(self, other: RunInfo) -> None:
        if not self.can_join(other):
            raise RuntimeError("Cannot join runs")

        self.text += other.text


def join_runs(runs: list[RunInfo]) -> list[RunInfo]:
    if len(runs) <= 1:
        return runs

    joint_runs: list[RunInfo] = [runs[0]]
    for run in runs[1:]:
        prev_run = joint_runs[-1]
        if prev_run.can_join(run):
            prev_run.join(run)
        else:
            joint_runs.append(run)

    return joint_runs


def trim_paragraphs(paragraphs: list[str]) -> list[str]:
    if paragraphs == []:
        return []

    left = 0
    while left < len(paragraphs) and paragraphs[left] == "<br>":
        left += 1

    right = len(paragraphs)
    while right >= 0 and paragraphs[right - 1] == "<br>":
        right -= 1

    if left >= right:
        return []

    return paragraphs[left:right + 1]


def is_uc_row(row: Row) -> bool:
    val: str | None = None
    for cell in row.cells:
        if val is None:
            val = cell.text
            continue

        if val != cell.text:
            return False

    return True


def is_empty_row(row: Row) -> bool:
    for cell in row.cells:
        if cell.text != "":
            return False

    return True


def ensure_json_file(file_name: str) -> str:
    _, ext = os.path.splitext(file_name)
    if ext != ".json":
        file_name += ".json"

    return file_name


def sanitize_text(text: str) -> str:
    return (
        text
        .replace("\u2012", "-")
        .replace("\u2013", "-")
        .replace("\n", "")
    )


class DocxParser:
    def __init__(self, path: str) -> None:
        self._doc = Document(path)
        self._para_styles = self._get_styles(WD_STYLE_TYPE.PARAGRAPH)

    @property
    def tables(self) -> list[Table]:
        return self._doc.tables

    def _get_styles(self, style_type: BaseStyle) -> dict[str, BaseStyle]:
        return {
            style.name: style
            for style
            in self._doc.styles
            if style.type == style_type and style.name is not None
        }

    def get_paragraph_style(self, name: str) -> ParagraphStyle | None:
        return self._para_styles.get(name)

    def convert_to_html(self, paragraph: Paragraph) -> str:
        """Converts the `Paragraph` object into HTML."""

        runs = [RunInfo(run) for run in paragraph.runs]
        _html = "".join([run.html for run in join_runs(runs)])
        bullet_style = self.get_paragraph_style("List Paragraph")
        if bullet_style is None:
            raise RuntimeError("Unable to find a bullet list style")

        if paragraph.style == bullet_style:
            _html = f"<li>{_html}</li>"

        return _html

    def parse_sections(self, file_name: str = "section_descriptions.json") -> None:
        table = self.tables[0]
        if table.rows == []:
            return {}

        is_bulleted = False
        data = {}
        headers = [cell.text.strip() for cell in table.rows[0].cells[1:]]
        for row in table.rows[1:]:
            row_data = {}
            for i, cell in enumerate(row.cells[1:]):
                paragraphs: list[str] = []
                for paragraph in cell.paragraphs:
                    if paragraph.text == "":
                        paragraphs.append("<br>")
                        continue

                    _html = self.convert_to_html(paragraph)
                    if not _html.startswith("<li>") and is_bulleted:
                        paragraphs[-1] += "</ul>"
                        is_bulleted = False

                    if _html.startswith("<li>") and not is_bulleted:
                        is_bulleted = True
                        _html = "<ul>" + _html

                    paragraphs.append(_html)

                if is_bulleted:
                    paragraphs[-1] += "</ul>"
                    is_bulleted = False

                row_data[headers[i]] = "".join(trim_paragraphs(paragraphs))

            data[row.cells[0].text] = row_data

        file_name = ensure_json_file(file_name)
        with open(file_name, "w+") as fp:
            json.dump(data, fp, indent=4)

    def parse_sunsetted_measures(self, file_name: str = "sunsetted_measures.json") -> None:
        table = self.tables[0]
        if table.rows == []:
            raise RuntimeError("Table is empty")

        data: dict[str, list] = {}
        cur_uc: str | None = None
        for row in table.rows[1:]:
            if is_empty_row(row):
                continue

            if is_uc_row(row):
                cur_uc = sanitize_text(row.cells[0].text)
                data[cur_uc] = []
                continue

            data[cur_uc].append(
                {
                    "version_id": sanitize_text(row.cells[1].text),
                    "name": sanitize_text(row.cells[2].text),
                    "active_life": sanitize_text(row.cells[3].text),
                    "trm_update": row.cells[4].text.upper() == "X",
                    "version_update": row.cells[5].text.upper() == "X",
                    "is_sunsetted": row.cells[6].text.upper() == "X"
                }
            )

        json_data: dict[str, list] = {
            "use_categories": []
        }
        for use_category, measures in data.items():
            json_data["use_categories"].append(
                {
                    "use_category": use_category,
                    "measures": measures
                }
            )

        file_name = ensure_json_file(file_name)
        with open(file_name, "w+") as fp:
            json.dump(json_data, fp, indent=4)


def main() -> None:
    parser = DocxParser("section_descs.docx")
    parser.parse_sections()


if __name__ == "__main__":
    main()
